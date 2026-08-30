"""Build a downloadable Word summary of a tender readiness checklist.

Produces a .docx that opens natively in Word, Google Docs, and LibreOffice, so
the user can download it, or upload it to Google Docs and keep working there.

Everything the checker found goes in, and everything it could NOT find goes in
too, under its own heading - a bidder who reads only this document still learns
what they have to confirm on the official page.
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x2E)
MUTED = RGBColor(0x5A, 0x5F, 0x66)
FLAG = RGBColor(0x9A, 0x54, 0x12)

CONTENT_WIDTH = Inches(6.5)          # Letter minus 1" margins
COLUMN_WIDTHS = (Inches(1.9), Inches(3.3), Inches(1.3))


def _shade(cell, hex_fill: str) -> None:
    """Cell background. Uses a clear pattern - 'solid' renders black in Word."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell(cell, text: str, width, bold: bool = False, color: RGBColor | None = None,
              size: int = 10) -> None:
    cell.width = width               # cells need their own width, not just the table
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _rule(document) -> None:
    """Horizontal rule as a paragraph bottom border, never as a 1-cell table."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(10)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C8CCC9")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _as_text(value) -> str:
    if isinstance(value, list):
        return ", ".join(str(item) for item in value)
    return str(value if value not in (None, "") else "-")


def build(checklist: dict, attachments: list | None = None, summary_override: str = "",
          language_label: str = "English") -> io.BytesIO:
    """Render the checklist to a .docx in memory. Returns a seeked BytesIO."""
    document = Document()

    section = document.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)   # US Letter
    for attribute in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attribute, Inches(1))

    fields = checklist.get("fields", {})
    title = fields.get("tender_title", {}).get("value") or "Tender notice"
    authority = fields.get("issuing_authority", {}).get("value") or "Issuing department not published"

    # --- header -----------------------------------------------------------
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    brand = heading.add_run("TENDERSATHI  ·  TENDER READINESS SUMMARY")
    brand.bold = True
    brand.font.size = Pt(8.5)
    brand.font.color.rgb = MUTED

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(4)
    title_run = title_paragraph.add_run(_as_text(title))
    title_run.bold = True
    title_run.font.size = Pt(17)
    title_run.font.color.rgb = ACCENT

    authority_paragraph = document.add_paragraph()
    authority_paragraph.paragraph_format.space_after = Pt(10)
    authority_run = authority_paragraph.add_run(_as_text(authority))
    authority_run.font.size = Pt(11)
    authority_run.font.color.rgb = MUTED

    found = checklist.get("found_count", 0)
    total = checklist.get("total_count", 0)
    meta_rows = [
        ("Tender ID", _as_text(fields.get("tender_id", {}).get("value"))),
        ("Extraction status", f"{found} of {total} fields extracted from the notice"),
        ("Summary language", language_label),
        ("Generated", datetime.now().strftime("%d %b %Y, %I:%M %p")),
    ]
    if checklist.get("source_url"):
        meta_rows.append(("Official source", checklist["source_url"]))

    meta_table = document.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in meta_rows:
        row = meta_table.add_row()
        _set_cell(row.cells[0], label, Inches(1.9), bold=True, color=MUTED, size=9)
        _set_cell(row.cells[1], value, Inches(4.6), size=9)

    _rule(document)

    # --- plain-language summary -------------------------------------------
    document.add_heading("Summary", level=1)
    summary_paragraph = document.add_paragraph(summary_override or checklist.get("summary", ""))
    summary_paragraph.paragraph_format.space_after = Pt(10)

    # --- requirements and details ------------------------------------------
    document.add_heading("Requirements and details", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # Column widths must be set on the table AND every cell, in absolute units.
    for index, width in enumerate(COLUMN_WIDTHS):
        table.columns[index].width = width

    header_cells = table.rows[0].cells
    for index, (label, width) in enumerate(zip(("Item", "Details", "Source"), COLUMN_WIDTHS)):
        _set_cell(header_cells[index], label, width, bold=True, size=9)
        _shade(header_cells[index], "EDF1EE")

    for field in fields.values():
        row = table.add_row().cells
        _set_cell(row[0], field.get("label", ""), COLUMN_WIDTHS[0], bold=True, size=9.5)
        if field.get("found"):
            _set_cell(row[1], _as_text(field.get("value")), COLUMN_WIDTHS[1], size=9.5)
            _set_cell(row[2], field.get("source", ""), COLUMN_WIDTHS[2], color=MUTED, size=8.5)
        else:
            _set_cell(row[1], "Not published in the notice", COLUMN_WIDTHS[1], color=FLAG, size=9.5)
            _set_cell(row[2], "confirm on official page", COLUMN_WIDTHS[2], color=FLAG, size=8.5)
            for cell in row:
                _shade(cell, "FDF6EE")

    document.add_paragraph()

    # --- the part that decides whether a bid is valid ----------------------
    missing = checklist.get("missing", [])
    document.add_heading("Before you bid", level=1)
    if missing:
        lead = document.add_paragraph()
        lead_run = lead.add_run(
            f"{len(missing)} of {total} items were not published in the notice that was read. "
            "Confirm each one on the official tender page before submitting - a missing "
            "eligibility line or document is enough to disqualify a bid."
        )
        lead_run.font.color.rgb = FLAG
        for item in missing:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        document.add_paragraph(
            "Every field was extracted from the notice. Still verify the values on the "
            "official tender page before submitting."
        )

    for warning in checklist.get("warnings", []):
        note = document.add_paragraph()
        note_run = note.add_run("Note: " + str(warning))
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = MUTED

    # --- attached document images ------------------------------------------
    if attachments:
        document.add_paragraph().add_run().add_break()
        document.add_heading("Attached documents", level=1)
        document.add_paragraph(
            "Pages of the tender document as uploaded. These are the source pages the "
            "fields above were read from."
        )
        for index, attachment in enumerate(attachments, start=1):
            caption = document.add_paragraph()
            caption.paragraph_format.space_before = Pt(10)
            caption.paragraph_format.space_after = Pt(4)
            caption_run = caption.add_run(attachment.get("caption") or f"Page {index}")
            caption_run.bold = True
            caption_run.font.size = Pt(9)
            caption_run.font.color.rgb = MUTED
            try:
                document.add_picture(attachment["path"], width=CONTENT_WIDTH)
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                failed = document.add_paragraph()
                failed_run = failed.add_run("This page could not be embedded.")
                failed_run.italic = True
                failed_run.font.color.rgb = FLAG

    # --- disclaimer ---------------------------------------------------------
    _rule(document)
    disclaimer = document.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "Guidance only. TenderSathi summarises a tender notice to help you prepare; it does "
        "not submit bids and is not legal advice. The official tender page remains the source "
        "of truth for every requirement, value, and deadline."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(8.5)
    disclaimer_run.font.color.rgb = MUTED

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def filename_for(checklist: dict) -> str:
    """A filename that reads well in a downloads folder."""
    fields = checklist.get("fields", {})
    stem = fields.get("tender_id", {}).get("value") or fields.get("tender_title", {}).get("value") or "tender"
    safe = "".join(character if character.isalnum() or character in "-_" else "-" for character in str(stem))
    return f"tender-readiness-{safe.strip('-')[:60] or 'summary'}.docx"
